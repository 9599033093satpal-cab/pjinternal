"""
Aether OCR — Legal Draft Generator
=====================================
Generates court-ready DOCX and PDF from master_case.json.
Uses python-docx template placeholders.
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Run: pip install python-docx")

try:
    import docx2pdf
    PDF_CONVERT = True
except ImportError:
    PDF_CONVERT = False


TEMPLATES_DIR = Path(__file__).parent / "legal_templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


class LegalDraftGenerator:
    """Generates legal documents from master_case.json."""

    SUPPORTED_DRAFTS = {
        "sarfaesi_demand_notice": "SARFAESI Section 13(2) Demand Notice",
        "possession_notice":      "SARFAESI Section 13(4) Possession Notice",
        "loan_recall_letter":     "Loan Recall Letter",
        "legal_opinion_report":   "Legal Opinion Report",
        "auction_notice":         "E-Auction Notice",
    }

    def generate(self, master_case: dict, draft_type: str, output_dir: str) -> dict:
        """
        Generate a legal draft.
        Returns: {success, docx_path, pdf_path, draft_type, generated_at}
        """
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}

        if draft_type not in self.SUPPORTED_DRAFTS:
            return {"success": False, "error": f"Unknown draft type: {draft_type}"}

        try:
            generator_fn = getattr(self, f"_gen_{draft_type}", None)
            if generator_fn:
                doc = generator_fn(master_case)
            else:
                doc = self._gen_generic(master_case, draft_type)

            # Save DOCX
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            docx_filename = f"{draft_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            docx_path = output_path / docx_filename
            doc.save(str(docx_path))

            # Convert to PDF if available
            pdf_path = None
            if PDF_CONVERT:
                try:
                    pdf_path = str(docx_path).replace(".docx", ".pdf")
                    docx2pdf.convert(str(docx_path), pdf_path)
                except Exception as e:
                    logger.warning(f"PDF conversion failed: {e}")
                    pdf_path = None

            logger.info(f"Draft generated: {docx_path}")
            return {
                "success": True,
                "docx_path": str(docx_path),
                "pdf_path": pdf_path,
                "draft_type": draft_type,
                "draft_name": self.SUPPORTED_DRAFTS[draft_type],
                "generated_at": datetime.now().isoformat(),
            }

        except Exception as e:
            logger.error(f"Draft generation failed: {e}")
            return {"success": False, "error": str(e)}

    def _gen_sarfaesi_demand_notice(self, case: dict) -> "Document":
        """Generate SARFAESI Section 13(2) Demand Notice."""
        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(case.get("parties", {}).get("lender", {}).get("bank_name", "BANK NAME"))
        run.bold = True
        run.font.size = Pt(14)

        branch = case.get("parties", {}).get("lender", {}).get("branch", "Branch Name")
        doc.add_paragraph(branch).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Notice title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t = title.add_run("NOTICE UNDER SECTION 13(2) OF THE SECURITISATION AND RECONSTRUCTION OF\nFINANCIAL ASSETS AND ENFORCEMENT OF SECURITY INTEREST ACT, 2002")
        t.bold = True
        t.font.size = Pt(11)

        doc.add_paragraph()

        # Date
        notice_date = case.get("legal_status", {}).get("notice_date", datetime.now().strftime("%d/%m/%Y"))
        doc.add_paragraph(f"Date: {notice_date}")
        doc.add_paragraph()

        # Addressee
        borrower = case.get("parties", {}).get("borrower", {})
        doc.add_paragraph("To,")
        doc.add_paragraph(borrower.get("name", "[BORROWER NAME]"))
        doc.add_paragraph(borrower.get("address", "[BORROWER ADDRESS]"))
        doc.add_paragraph()

        co_borrower = case.get("parties", {}).get("co_borrower", {})
        if co_borrower.get("name"):
            doc.add_paragraph("And,")
            doc.add_paragraph(co_borrower["name"])
            doc.add_paragraph()

        # Subject
        subj = doc.add_paragraph()
        subj.add_run("Sub: ").bold = True
        financial = case.get("financial", {})
        subj.add_run(
            f"Demand Notice under Section 13(2) of SARFAESI Act, 2002 in respect of "
            f"Loan Account No. {financial.get('loan_number', '[LOAN NO]')} "
            f"— Outstanding Amount of ₹{financial.get('total_demand', '[AMOUNT]')}"
        )

        doc.add_paragraph()
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph()

        # Body
        body_text = (
            f"You had availed a loan/credit facility from {borrower.get('name', '[BANK NAME]')} "
            f"vide Loan Account No. {financial.get('loan_number', '[LOAN NO]')} "
            f"for an amount of ₹{financial.get('loan_amount', '[AMOUNT]')}. "
            f"As on date, the outstanding dues payable by you are as under:\n\n"
            f"  Principal Outstanding: ₹{financial.get('outstanding_principal', '[AMOUNT]')}\n"
            f"  Interest Outstanding:  ₹{financial.get('outstanding_interest', '[AMOUNT]')}\n"
            f"  Total Amount Payable:  ₹{financial.get('total_demand', '[AMOUNT]')}\n\n"
            f"The above loan account has been classified as a Non-Performing Asset (NPA) "
            f"in accordance with the guidelines of the Reserve Bank of India. "
            f"You are hereby called upon to pay the aforesaid total outstanding amount of "
            f"₹{financial.get('total_demand', '[AMOUNT]')} within 60 (sixty) days from the "
            f"date of receipt of this notice, failing which the Bank shall be constrained to "
            f"enforce the security interest created over the secured assets described hereunder, "
            f"without any further notice to you, as per the provisions of the SARFAESI Act, 2002."
        )
        doc.add_paragraph(body_text)
        doc.add_paragraph()

        # Property details
        prop = case.get("property", {})
        if any(prop.values()):
            prop_heading = doc.add_paragraph()
            prop_heading.add_run("DESCRIPTION OF SECURED ASSETS:").bold = True
            prop_text = (
                f"All that piece and parcel of property bearing Survey No. {prop.get('survey_number', '[SURVEY NO]')}, "
                f"admeasuring {prop.get('area', '[AREA]')} {prop.get('area_unit', 'sq.ft.')}, "
                f"situated at {prop.get('address', '[ADDRESS]')}, "
                f"District: {prop.get('district', '[DISTRICT]')}, "
                f"State: {prop.get('state', '[STATE]')}."
            )
            doc.add_paragraph(prop_text)
            doc.add_paragraph()

        # Closing
        doc.add_paragraph(
            "Please note that this notice is without prejudice to any other rights and remedies "
            "available to the Bank under law and/or under the loan documents executed by you."
        )
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        officer = case.get("parties", {}).get("lender", {}).get("officer", "[AUTHORIZED OFFICER]")
        doc.add_paragraph(f"{officer}")
        doc.add_paragraph("Authorized Officer")
        lender = case.get("parties", {}).get("lender", {})
        doc.add_paragraph(f"{lender.get('bank_name', '')}, {lender.get('branch', '')}")

        return doc

    def _gen_loan_recall_letter(self, case: dict) -> "Document":
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        doc.add_paragraph("LOAN RECALL LETTER").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph()

        borrower = case.get("parties", {}).get("borrower", {})
        financial = case.get("financial", {})
        lender = case.get("parties", {}).get("lender", {})

        doc.add_paragraph(f"To,\n{borrower.get('name', '[NAME]')}\n{borrower.get('address', '[ADDRESS]')}")
        doc.add_paragraph()
        doc.add_paragraph(
            f"Dear {borrower.get('name', 'Sir/Madam')},\n\n"
            f"We refer to the loan/credit facility granted to you vide "
            f"Loan Account No. {financial.get('loan_number', '[LOAN NO]')}. "
            f"Despite repeated reminders, you have failed to regularize your loan account. "
            f"We hereby recall the entire loan outstanding amount of "
            f"₹{financial.get('total_demand', '[AMOUNT]')} and demand payment immediately.\n\n"
            f"Failure to comply shall result in legal action without further notice."
        )
        doc.add_paragraph()
        doc.add_paragraph(f"Yours faithfully,\n\n\n{lender.get('officer', '[OFFICER]')}\nAuthorized Officer\n{lender.get('bank_name', '')}")
        return doc

    def _gen_generic(self, case: dict, draft_type: str) -> "Document":
        """Generic template for unsupported types."""
        doc = Document()
        doc.add_paragraph(self.SUPPORTED_DRAFTS.get(draft_type, draft_type)).runs[0].bold = True
        doc.add_paragraph(json.dumps(case.get("financial", {}), indent=2))
        return doc
